from docx import Document
from datetime import datetime
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH





def generate_meeting_report(template_path, output_path, meeting_details):
    """
    Generate a meeting report from a template
    
    Args:
        template_path: Path to the docx template
        output_path: Where to save the generated report
        meeting_description: Brief description of the meeting
        meeting_details: Dict containing meeting metadata (date, attendees, etc.)
    """
    # Load the template
    doc = Document(template_path)
    
    # Replace placeholders in the document
    # for paragraph in doc.paragraphs:
    #     if '[ATTENDEES]' in paragraph.text:
    #         paragraph.text = paragraph.text.replace('[ATTENDEES]', ', '.join(meeting_details.get('attendees', [])))
            
    #     if '[MEETING_DESCRIPTION]' in paragraph.text:
    #         paragraph.text = paragraph.text.replace('[MEETING_DESCRIPTION]', meeting_description)
    
    # sections = doc.sections
    # print(len(sections))

    for paragraph in doc.paragraphs:
        for key, value in meeting_details.items():
            if f'[{key}]' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'[{key}]', str(value))
    
    # Handle tables if needed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in meeting_details.items():
                    if f'[{key}]' in cell.text:
                        cell.text = cell.text.replace(f'[{key}]', str(value))

    for section in doc.sections: 
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in meeting_details.items():
                        if f'[{key}]' in cell.text:
                            cell.text = cell.text.replace(f'[{key}]', str(value))
                            if f'[{key}]' == "[TITLE]":
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER 

        # section.header.tables[0].cell(2,1).paragraphs[0].text = section.header.tables[0].cell(2,1).paragraphs[0].text.replace("[TITLE]",meeting_details.get('title'))


    
    # Save the modified document
    doc.save(output_path)
    print(f"Report generated at {output_path}")

# Example usage
meeting_details = {
    'MEETING_DATE': '2023-11-15',
    'TITLE': 'Quarterly Planning Session',
    'MEETING_DESCRIPTION': "We discussed the Q4 goals and allocated resources for upcoming projects."
}

generate_meeting_report(
    template_path='meeting_template.docx',
    output_path='generated_report.docx',
    meeting_details=meeting_details
)

