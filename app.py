from docx import Document
from datetime import datetime

def generate_meeting_report(template_path, output_path, meeting_description, meeting_details):
    """
    Generate a meeting report from a template
    
    Args:
        template_path: Path to the docx template
        output_path: Where to save the generated report
        meeting_description: Brief description of the meeting
        meeting_details: Dict containing meeting metadata (date, attendees, etc.)
    """
    # Load the template
    doc = Document(template_path)
    
    # Replace placeholders in the document
    for paragraph in doc.paragraphs:
        if '[ATTENDEES]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[ATTENDEES]', ', '.join(meeting_details.get('attendees', [])))
            
        if '[MEETING_DESCRIPTION]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[MEETING_DESCRIPTION]', meeting_description)
    
    sections = doc.sections
    print(len(sections))

    for section in sections: 
        print(section.header.tables[0].cell(2,1).paragraphs[0].text)
        section.header.tables[0].cell(2,1).paragraphs[0].text = section.header.tables[0].cell(2,1).paragraphs[0].text.replace("[test]","hello")

    # sections[0].header.paragraphs[0].text = sections[0].header.paragraphs[0].text.replace("[test]","hello")
    
    # Save the modified document
    doc.save(output_path)
    print(f"Report generated at {output_path}")

# Example usage
meeting_details = {
    'date': '2023-11-15',
    'title': 'Quarterly Planning Session',
    'attendees': ['John Doe', 'Jane Smith', 'Bob Johnson']
}

generate_meeting_report(
    template_path='meeting_template.docx',
    output_path='generated_report.docx',
    meeting_description='We discussed the Q4 goals and allocated resources for upcoming projects.',
    meeting_details=meeting_details
)